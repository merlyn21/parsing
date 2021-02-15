#!/usr/bin/python3
# -*- coding: utf-8 -*-

import requests
import codecs
import openpyxl
from bs4 import BeautifulSoup
import re
import docx
import time
import datetime
import urllib.parse
import zipfile
import os
import glob

import searchtext
import sendfile

from openpyxl.styles import PatternFill, Border, Side, Alignment, Protection, Font, Side
# from datetime import  *


month_list = ['январь', 'февраль', 'март', 'апрель', 'май', 'июнь', 'июль', 'август', 'сентябрь', 'октябрь', 'ноябрь', 'декабрь']

path_scr = os.path.dirname(os.path.realpath(__file__))

font = Font(name='Calibri',
            size=12,
            bold=True,
            italic=False,
            vertAlign=None,
            underline='none',
            strike=False,
            color='FF000000')

now = datetime.datetime.now()
mec = month_list[int(now.strftime("%m")) - 1]
day = now.strftime("%d")
file_suf = now.strftime("%d-%m-%Y_%H-%M")
filename2send = path_scr+'/reports/query.docx'
count_cell = 0

emailto = 'merlyn022@yandex.ru'
subject = 'Рeports word '+file_suf
message = 'Смотрите вложение'

soderjanie = ' '
count_s = 1

#str_find = 'вчера'.decode('utf-8')

def zipdir(path, ziph):
    # ziph is zipfile handle
    for root, dirs, files in os.walk(path):
        for file in files:
            ziph.write(os.path.join(root, file))



def get_query(soup, doc, name_q):

    global count_cell
    len_m=len(soup.find_all(class_='document__provider'))

    print('From get_query. Длина '+str(len_m))
    i = 0
    j = 0
    while i < len_m:
        # print(soup.find_all(class_='document__provider')[i].get_text())
        # print(soup.find_all(class_='document__time')[i].get_text())
        # print(soup.find_all(class_='document__title')[i].get_text())
        # print(soup.find_all(class_='document__title')[i].a["href"])
        gettext = soup.find_all(class_='document__snippet')[i].get_text()
        # print(soup.find_all(class_='document__snippet')[i].get_text())

        dt = soup.find_all(class_='document__time')[i].get_text()
        # print(re.fullmatch('вчера в \d\d\D\d\d',dt))
        print("Data = " + soup.find_all(class_='document__time')[i].get_text())
        global soderjanie
        global count_s
        if re.fullmatch('\d\d\D\d\d', dt) or re.search('вчера', dt):
            #print('write to Excel cl = ' + str(count_cell))
            #cell = sheet.cell(count_cell+2, 1)
            #cell.value = mec
            #cell = sheet.cell(count_cell+2, 2)
            #cell.value = day
            #cell = sheet.cell(count_cell+2, 3)
            #cell.value = soup.find_all(class_='document__title')[i].get_text()
            href = soup.find_all(class_='document__title')[i].a["href"]
            #cell.hyperlink = href
            #cell.style = 'Hyperlink'

            #cell = sheet.cell(count_cell+2, 4)
            #cell.value = soup.find_all(class_='document__provider')[i].get_text()
            #cell = sheet.cell(count_cell+2, 5)
            #cell.value = soup.find_all(class_='document__time')[i].get_text()
            #cell = sheet.cell(count_cell + 2, 6)
            #cell.value = name_q

            print('Find for word! ++++++++++++++++++++++++++++++++++++++++++++++')
            print(href)
            print(gettext)
            print(soup.find_all(class_='document__time')[i].get_text())
            print('++++++')
            #doc.add_paragraph("------------------------------------------------------- \n\n")
            zag = doc.add_paragraph().add_run(soup.find_all(class_='document__title')[i].get_text())
            zag.bold = True
            doc.add_paragraph(soup.find_all(class_='document__time')[i].get_text()+'   '+soup.find_all(class_='document__provider')[i].get_text())
            doc.add_paragraph(searchtext.get_text(href, gettext[5:]))
            doc.add_paragraph('URL: '+ href)
            doc.add_paragraph("------------------------------------------------------- \n\n")

            soderjanie = soderjanie + str(count_s)+ '. ' + soup.find_all(class_='document__title')[i].get_text() + '\n'
            count_s+=1

            j+=1
            count_cell+=1

        i+=1
    #file.write(now.isoformat() + ": " + name_q + " - Обработали " + str(i) + " запросов, из них подошло " + str(j) + "\n")
    # with open('test.html', 'w') as output_file:
    #     output_file.write(soup.prettify())



def get_soup(query, pages, geonews):
    url = 'https://news.yandex.ru/yandsearch'

    print('Start get_soup! ' + str(pages))

    payload = {'text': query,
               'rpt': 'nnews2',
              # 'grhow': 'clutop',
              #  'wiz_no_news': '1',
               'rel': 'tm',
               'within': '8'}

    referer = url + '?' + payload["text"] + '&rpt=' + payload["rpt"] + '&rel=' + payload["rel"] + '&within=' + payload["within"]
    print(referer)
    if pages > 0:
        payload["p"] = pages
        print(payload)
        referer = referer+'&p='+str(payload["p"])

    if geonews > 0:
        payload["geonews"] = geonews
        referer = referer + '&geonews=' + str(payload["geonews"])



    headers = {
            'Referer': urllib.parse.quote(referer),
            #'Refer': 'https://yandex.ru/',
            'Connection': 'keep-alive',
            'Origin': 'https://news.yandex.ru',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/77.0.3865.90 Safari/537.36',
            'Sec-Fetch-Mode': 'cors',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3',
            'Accept-Encoding': 'gzip, deflate, br',
            'Accept-Language': 'ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7',
            'Cache-Control': 'max-age=0',
            'Host': 'news.yandex.ru',
            'Sec-Fetch-Mode': 'navigate',
            'Sec-Fetch-Site': 'same-site',
            'Sec-Fetch-User': '?1',
            'Upgrade-Insecure-Requests': '1',
            'Cookie': 'yandexuid=5026847811569230343; _ym_uid=1569230355874526059; mda=0; my=YwA=; mynews=0%3A1; fuid01=5dc174c61e7b817b.ANy-bTpxf83uZroLhodQ8TTCWDZuh9H-ERyYReBhyeJdxmi1H89j458zWpWeXAXRl65DM6YF6S4nahU0PLiOgZ1_jYX_PIfhpe1CcGUiDLJVTBZVkSM8IxKoqk5hV975; yandex_gid=213; _ym_d=1573135860; zm=m-white_bender.webp.css-https%3As3home-static_oPOgCkS33PyFVY9YQmBT1UvFErs%3Al; yc=1573395062.zen.cach%3A1573139459; L=cyFJSVBwYQNKAmUBDQlBf3lnUVhBbUpzOjoAPEIqOhQfECEkACR8EA44.1573135875.14042.315467.ea7098c93e96539cd89b0070d92d7652; yandex_login=robot@intelkon.com; Cookie_check=1; i=y25EZExBZ17zj9C2h2jmbGedHlUXbbOKMC2QLS9N3xMl9qZTmeh9GPtoyCYyFu2F+k1UOelbvM6h5jVK6Tzs3h5d0/g=; Session_id=3:1573655113.5.0.1573135875194:YvPsTQ:83.1|1130000041464657.0.2|207997.933037.0zJZFhPeXroN7nYZ84uZ330UHmU; sessionid2=3:1573655113.5.0.1573135875194:YvPsTQ:83.1|1130000041464657.0.2|207997.326626.ZhO5zXTX687CqHWzN_w8SPia0qM; _ym_isad=2; yp=1605351728.p_sw.1573815727#1585000576.szm.1_75:1920x1080:1097x554#1605351735.p_cl.1573815734#1605096742.wzrd_sw.1573560741#1605096742.dsws.29#1605096742.dswa.0#1605096742.dwsets.29#1888495859.multib.1#1575727859.ygu.1#1888495875.udn.cDpyb2JvdEBpbnRlbGtvbi5jb20%3D#1605352212.stltp.serp_bk-map_1_1573816212; ys=udn.cDpBbmRyZXlMZW9udGlldg%3D%3D#ymrefl.14C4B04EB1209767#wprid.1573829545753097-1212989197783622925500130-sas4-3060; _ym_visorc_722818=b; _ym_visorc_93511=w; yabs-frequency=/4/1m0E012_pbrXhyvT/vbMmS5Gu8GPKi71NEFp___zBLB1mL3WXFskmS5Su_F___mrMi71KE264LR1mL3XW-6cmS5Cu_F___mLKi70pEFp___yKLh1mL3X0zswmS2yu_F___-TLi71KE43KRh1m5pZy____wrMmS5Gu8C5gi71FEFp___yTRR1mL3Y02csmS5Gu8Effi71KE208RR1mL3WWZ5MmS5KuG0FxLB1m0pZy____2MsmSFCt_F___m00/; cycada=zEhXW3KkGHA6ifb7NSbWNo88sbI+2y+Qs9MK0NuX+F4='
          }


    r = requests.get(url, headers = headers, params=payload)

    if r.status_code == 200:
        print('Все в норме get_soup! ')
        r.encoding = 'utf-8'
        soup = BeautifulSoup(r.text, 'html.parser')
        return soup
    else:
        print(r.text)
        print(r.status_code)
        return -1


def get_main(query, doc, name_q, geonews=0):
    print("Start get_main!  " + query)
    doc.add_heading(name_q, 1)
    global soderjanie
    soderjanie = soderjanie + '\n' + name_q +'\n'
    file = open(path_scr+"/message.log", "a")



    soup = get_soup(query, 0, geonews)
    #print(soup)

    #now = datetime.datetime.now()

    if soup.find("title"):
        if soup.find("title").get_text() == "Ой!":
            print(now.isoformat() + ": - Попали в спам!")
            print(soup)
            file.write(now.isoformat() + ": " + name_q + " - Попали в спам!\n")
            return -1


    if soup:
        file.write(now.isoformat() + ": " + name_q + " - Суп получен, приступаем к работе!\n")
        get_query(soup, doc, name_q)
    else:
        print("Error get_soup")
        return -1

    i = 1
    # print(str(len(soup.find_all('span', {'class':'pager__group'}))))
    span = soup.find_all('span', {'class':'pager__group'})
    if span:
        print("two and more pages")
        #print(span[0])
        #print(str(len(span[0].find_all('a'))))
        count_pages = len(span[0].find_all('a'))
        while i < count_pages:
            #print(span[0].find_all('a')[i]['href'])
            soup = get_soup(query, i, geonews)
            get_query(soup, doc, name_q)
            i+=1

    else:
        print("one page!")


    # sheet.column_dimensions['A'].width = 15
    # sheet.column_dimensions['B'].width = 5
    # sheet.column_dimensions['C'].width = 80
    # sheet.column_dimensions['D'].width = 40
    # sheet.column_dimensions['E'].width = 20

    # if count_cell > 0:
    #     wb.save('reports\/query.xlsx')
    #     doc.save('reports\/query.docx')

    file.close()
    return 1


#wb = openpyxl.Workbook()
#wb.create_sheet(title='Первый лист', index=0)
#sheet = wb['Первый лист']

#cell = sheet.cell(1, 1)
#cell.value = 'Месяц'
#cell.font = font
#cell = sheet.cell(1, 2)
#cell.value = 'Дата'
#cell.font = font
#cell = sheet.cell(1, 3)
#cell.value = 'Публикация(Заголовок)'
#cell.font = font
#cell = sheet.cell(1, 4)
#cell.value = 'Название СМИ'
#cell.font = font
#cell = sheet.cell(1, 5)
#cell.value = 'Время'
#cell.font = font

doc = docx.Document()

sod = doc.add_paragraph()
sod.bold = True


#url1 = 'https://news.yandex.ru/yandsearch'
yuzhd = '(юужд | южно-уральская железная | Вокзал Челябинск | Вокзал Курган | Вокзал Оренбург | Вокзал Петропавловск | поезд Челябинск | поезд Курган | поезд Оренбург | поезд Петропавловск | электричка Челябинск | электричка Курган | электричка Оренбург | электричка Петропавловск | железная дорога Челябинск | железная дорога Курган | железная дорога Оренбург | южноуральская железная)  -трамвай -Мамаев'
chel = '(Вокзал | поезд | тепловоз | электровоз | паровоз | электричка | железная дорога | железнодорожный | РЖД | вагон | под вагонами | цистерна | цистрены | южно-уральская дирекция | электропоезд | южноуральская железная | ЮУЖД | Южно-Уральская железная | Южноуральская железная) -трамвай; -автовокзал; -трамвайных'
kurgan = '(Вокзал | поезд | тепловоз | электровоз | паровоз | электричка | железная дорога | железнодорожный | ржд | вагон | под вагонами | цистерна | цистерны | электропоезд | ЮУЖД | Южно-Уральская железная | Южноуральская железная | южно-уральская дирекция | анатолий храмцов | южноуральская железная) -трамвай -трамвайных -автовокзал'
orenburg ='(Вокзал | поезд | тепловоз | электровоз | паровоз | электричка | электропоезд | железная дорога | железнодорожный | ржд | вагон | под вагонами | цистерна | цистерн | ЮУЖД | южно-уральская дирекция | Южно-Уральская железная | Южноуральская железная | южноуральская железная) -трамвай -трамвайных -автоцистерна'
regions = '(Вокзал Челябинская | Вокзал Курганская | Вокзал Оренбургская | Вокзал Северо-Казахстанская | поезд Челябинская | поезд Курганская | поезд Оренбургская | поезд Северо-Казахстанская | электричка Челябинская | электричка Курганская | электричка Оренбургская | электричка Северо-Казахстанская | железная дорога Челябинская | железная дорога Курганская | железная дорога Оренбургская) -трамвай'
#query1 = 'юужд | южно-уральская железная | Вокзал Челябинск'


files = glob.glob(path_scr+'/reports/*')
for f in files:
    os.remove(f)

#files = glob.glob(path_scr+'/query*.zip')
#for f in files:
#    os.remove(f)

get_main(yuzhd, doc, 'ЮУЖД', )
time.sleep(35)
get_main(chel, doc, 'Челябинск', 56)
time.sleep(45)
get_main(kurgan, doc, 'Курган', 53)
time.sleep(55)
get_main(orenburg, doc, 'Оренбург', 48)
time.sleep(62)
get_main(regions, doc, 'Регионы')

#sheet.column_dimensions['A'].width = 15
#sheet.column_dimensions['B'].width = 5
#sheet.column_dimensions['C'].width = 80
#sheet.column_dimensions['D'].width = 40
#sheet.column_dimensions['E'].width = 20

print(path_scr)
#count_cell = 1
if count_cell > 0:
#    wb.save(path_scr+'/reports/query.xlsx')
    soderjanie = soderjanie+'\n\n\n\n----------------------------------------\n'
    sod.add_run(soderjanie)
    doc.save(path_scr+'/reports/query.docx')

#    zf = zipfile.ZipFile(filename2send, "w")
#    for dirname, subdirs, files in os.walk(path_scr+"/reports/"):
#        for filename in files:
#            zf.write(os.path.relpath(os.path.join(dirname, filename)))
#    zf.close()

    sendfile.send2mail(emailto, subject, message, filename2send)


#get_query(url1, query1)
